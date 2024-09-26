import os
from flask import Blueprint, render_template, request, jsonify, send_file
import requests
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from fpdf import FPDF
from dotenv import load_dotenv

load_dotenv()

reports_bp = Blueprint('reports', __name__)

api_key = os.getenv("API_KEY")
api_url = os.getenv("API_URL")

def clean_generated_content(content):
    unwanted_keywords = ["Title:", "Report:"]
    lines = content.split('\n')
    cleaned_lines = [line for line in lines if not any(line.startswith(keyword) for keyword in unwanted_keywords)]
    return '\n'.join(cleaned_lines)

def generate_report(topic, tone, length, industry, audience):
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    
    prompt = f"Generate a {length}-page {tone} report on {topic}, targeting {audience} in the {industry} industry."

    payload = {
        "model": "tiiuae/falcon-180B-chat",
        "messages": [
            {"role": "system", "content": "Generate a report based on user input."},
            {"role": "user", "content": prompt}
        ],
        "max_tokens": 1000
    }

    try:
        response = requests.post(api_url, headers=headers, json=payload)
        response.raise_for_status()

        data = response.json()
        report_content = data["choices"][0]["message"]["content"]
        return clean_generated_content(report_content)

    except requests.exceptions.RequestException as e:
        raise Exception(f"Error in API request: {e}")

@reports_bp.route('/generate_report', methods=['POST'])
def generate_report_endpoint():
    data = request.json

    topic = data.get('instructions')
    tone = data.get('tone')
    length = data.get('length')
    industry = data.get('industry')
    audience = data.get('audience')

    try:
        generated_report = generate_report(topic, tone, length, industry, audience)

        # Return the generated report as JSON for frontend display (no download)
        return jsonify({'generated_report': generated_report})

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@reports_bp.route('/download_report', methods=['POST'])
def download_report():
    data = request.json
    content = data.get('content')
    format = data.get('format', 'docx')

    try:
        if format == 'pdf':
            return generate_pdf(content)
        else:
            return generate_docx(content)

    except Exception as e:
        return jsonify({'error': str(e)}), 500


def generate_docx(content):
    try:
        doc = Document()

        # Adding a styled title
        title = doc.add_heading('Generated Report', level=1)
        run = title.runs[0]
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0, 112, 192)  # Blue color for title
        run.bold = True
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add paragraphs with styling
        for paragraph in content.split('\n'):
            if paragraph.strip():
                para = doc.add_paragraph(paragraph)
                run = para.runs[0]
                run.font.size = Pt(12)
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Debugging: Save to local file to check if docx is generated correctly
        temp_file = "/tmp/generated_report.docx"
        doc.save(temp_file)

        # Return file from disk
        return send_file(temp_file, as_attachment=True, download_name='generated_report.docx')

    except Exception as e:
        print(f"Error generating DOCX: {e}")
        raise

def generate_pdf(content):
    try:
        pdf = FPDF()
        pdf.add_page()

        # Add styled title
        pdf.set_font("Arial", "B", 16)
        pdf.set_text_color(0, 112, 192)  # Blue color for title
        pdf.cell(200, 10, txt="Generated Report", ln=True, align='C')

        # Add paragraphs
        pdf.set_text_color(0, 0, 0)  # Black for content
        pdf.set_font("Arial", size=12)
        for line in content.split('\n'):
            pdf.multi_cell(0, 10, line)

        # Debugging: Save to local file to check if PDF is generated correctly
        temp_file = "/tmp/generated_report.pdf"
        pdf.output(temp_file)

        # Return file from disk
        return send_file(temp_file, as_attachment=True, download_name='generated_report.pdf')

    except Exception as e:
        print(f"Error generating PDF: {e}")
        raise
